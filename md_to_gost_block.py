# -*- coding: utf-8 -*-
"""
Пакетный конвертер Markdown -> DOCX по ГОСТ 7.32-2001 (новая версия)

Полный функционал из md_to_gost_docx_batch.py, но с улучшенной архитектурой:
- Модульная структура с разделением ответственности
- Легко читаемый и поддерживаемый код
- Полная поддержка всех элементов Markdown
- Корректные подписи к рисункам и таблицам
- Нумерация элементов
- Правильное форматирование по ГОСТ
"""

import os
import re
import subprocess
import tempfile
from copy import deepcopy
from typing import List, Optional, Dict, Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

from MarkToGost.config import DocumentSettings
from MarkToGost.parser.blocks import BaseBlock, FormulaBlock, Section, HeadingBlock, ImageBlock, TableBlock, ListBlock, \
    CodeBlock, TextBlock
from MarkToGost.parser.markdown_parser import MarkdownParser
from MarkToGost.parser.metadata import extract_metadata
from MarkToGost.utils.formatting import apply_italic_formatting, set_run_font, set_paragraph_formatting


def reset_heading_styles(doc: Document):
    """Сбрасывает форматирование стилей Heading 2-9 до нужного вида"""
    for level in range(2, 10):
        style_name = f'Heading {level}'
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue

        # Шрифт
        style.font.name = DocumentSettings.FONT_NAME
        style.font.size = Pt(DocumentSettings.FONT_SIZE_PT)
        style.font.bold = True
        style.font.italic = False
        style.font.all_caps = False
        style.font.small_caps = False
        style.font.color.rgb = RGBColor(0, 0, 0)

        # Абзац
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.left_indent = Cm(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = DocumentSettings.LINE_SPACING

        # Сброс caps через XML (на случай если python-docx не перебивает)
        rPr = style.element.get_or_add_rPr()
        for tag in ('w:caps', 'w:smallCaps'):
            el = rPr.find(qn(tag))
            if el is not None:
                rPr.remove(el)
        # Явно выставляем w:caps val=0
        caps_el = OxmlElement('w:caps')
        caps_el.set(qn('w:val'), '0')
        rPr.append(caps_el)





# ================================
# КОНФИГУРАЦИЯ И НАСТРОЙКИ
# ================================
def render_formula_with_pandoc(latex: str) -> Document:
    """Генерирует DOCX с формулой через Pandoc или текстовый резервный вариант"""

    with tempfile.TemporaryDirectory() as tmpdir:
        md_path = os.path.join(tmpdir, "formula.md")
        docx_path = os.path.join(tmpdir, "formula.docx")

        # ВАЖНО: display math
        md_content = f"$${latex}$$"
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)

        try:
            subprocess.run([
                "pandoc",
                md_path,
                "-o",
                docx_path
            ], check=True, capture_output=True, timeout=10)

            return Document(docx_path)
        except (FileNotFoundError, subprocess.TimeoutExpired, subprocess.CalledProcessError):
            # Если pandoc не установлен или произошла ошибка, создаем документ с текстом формулы
            doc = Document()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"${latex}$")
            run.font.size = Pt(DocumentSettings.FONT_SIZE_PT)
            return doc


def append_docx_content(src_doc: Document, dst_doc: Document):
    """Копирует содержимое одного docx в другой"""
    for element in src_doc.element.body:
        dst_doc.element.body.append(deepcopy(element))


def set_table_borders(table):
    """Установка границ таблицы через XML"""
    tbl = table._tbl
    tblPr = tbl.tblPr

    # Удаляем старые границы
    for child in list(tblPr):
        if child.tag.split('}')[-1] == 'tblBorders':
            tblPr.remove(child)

    tblBorders = OxmlElement('w:tblBorders')
    for border_type in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elm = OxmlElement(f'w:{border_type}')
        elm.set(qn('w:val'), 'single')
        elm.set(qn('w:sz'), '8')
        elm.set(qn('w:space'), '0')
        elm.set(qn('w:color'), '000000')
        tblBorders.append(elm)

    tblPr.append(tblBorders)


def set_repeat_table_header(row):
    """Установка повторения заголовка таблицы на каждой странице"""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


def add_page_number_centered(document):
    """Добавление номеров страниц в нижний колонтитул"""
    for section in document.sections:
        footer = section.footer
        for p in footer.paragraphs:
            p.clear()
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), 'PAGE')
        run = OxmlElement('w:r')
        run.append(fld)
        p._p.append(run)

        for r in p.runs:
            r.font.name = DocumentSettings.FONT_NAME
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)


def compute_image_width_cm(doc, fraction=0.70):
    """Вычисление ширины изображения"""
    sec = doc.sections[0]
    try:
        page_width_cm = sec.page_width.cm
    except:
        page_width_cm = float(sec.page_width) / 360000.0

    left = sec.left_margin.cm
    right = sec.right_margin.cm
    avail = page_width_cm - left - right
    return Cm(avail * fraction)


def replace_image_refs(text: str, image_refs: Dict[str, int]) -> str:
    """Замена ссылок на изображения @img_id на их номера"""
    if not text or not image_refs:
        return text

    result = text
    for img_id, fig_num in image_refs.items():
        result = result.replace(f"@{img_id}", f"рис. {fig_num}")

    return result


# ================================
# РАБОТА С ТАБЛИЦАМИ
# ================================

def split_md_table_row(line):
    """Разбор строки markdown-таблицы"""
    s = line.strip()
    if s.startswith('|'):
        s = s[1:]
    if s.endswith('|'):
        s = s[:-1]

    cells = []
    buf = []
    escape = False

    for ch in s:
        if escape:
            buf.append(ch)
            escape = False
            continue
        if ch == '\\':
            escape = True
            continue
        if ch == '|':
            cells.append(''.join(buf).strip())
            buf = []
            continue
        buf.append(ch)

    cells.append(''.join(buf).strip())
    return cells


def is_md_table_separator(line):
    """Проверка разделителя таблицы"""
    if '|' not in line:
        return False

    parts = split_md_table_row(line)
    if len(parts) < 2:
        return False

    for part in parts:
        p = part.replace(' ', '')
        if not re.fullmatch(r':?-{3,}:?', p):
            return False
    return True


def is_md_table_row(line):
    """Проверка строки таблицы"""
    s = line.strip()
    return bool(s) and ('|' in s)


CAPTION_HTML_RE = re.compile(r'^\s*<caption>\s*(.*?)\s*</caption>\s*$', re.I)
CAPTION_MD_RE = re.compile(r'^\s*(?:Таблица|Table)\s*\d*\s*[—:-]\s*(.+?)\s*$', re.I)
CAPTION_PLN_RE = re.compile(r'^\s*(?:Название таблицы|Caption)\s*[:\-]\s*(.+?)\s*$', re.I)


def normalize_table_caption(text):
    """Извлечение названия таблицы"""
    if not text:
        return None

    text = re.sub(r'\s+', ' ', text).strip().rstrip('.')

    for regex in [CAPTION_HTML_RE, CAPTION_MD_RE, CAPTION_PLN_RE]:
        m = regex.match(text)
        if m:
            title = m.group(1).strip().rstrip('.')
            return title or None

    return None


def format_table_cell(cell, *, is_header=False, size_pt=12):
    """Форматирование ячейки таблицы"""
    cell.text = cell.text

    for p in cell.paragraphs:
        p.style = 'Normal'  # Изменено с None на 'Normal'
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT

        if not p.runs:
            p.add_run("")

        # Применяем курсив
        p.clear()
        for part_text, is_italic in apply_italic_formatting(cell.text):
            run = p.add_run(part_text)
            set_run_font(run, size_pt=size_pt, bold=is_header, italic=is_italic)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


# ================================
# БЛОКИ ДОКУМЕНТА
# ================================

# ================================
# ПАРСЕР MARKDOWN
# ================================



# ================================
# РЕНДЕРЕР ДОКУМЕНТА
# ================================



def add_toc(document):
    """Добавление автоматического оглавления (TOC)"""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_formatting(paragraph, space_before=0, space_after=6)

    run = paragraph.add_run()

    # Начало поля
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    # Инструкция TOC
    instrText = OxmlElement('w:instrText')
    instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instrText.text = 'TOC \\o "1-4" \\h \\z \\u'
    run._r.append(instrText)

    # Разделитель
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)

    # Конец поля
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)


# ================================
# ОСНОВНАЯ ЛОГИКА
# ================================

def get_heading_level_from_number(text: str) -> int:
    """
    Определяет уровень заголовка по нумерации:
    1. → 1
    1.1 → 2
    1.1.1 → 3
    """
    match = re.match(r'^(\d+(?:\.\d+)*)', text.strip())
    if not match:
        return 1  # если нет номера — считаем верхним уровнем

    number = match.group(1)
    level = number.count('.') + 1
    # Защита: Word поддерживает максимум Heading 1-9
    return min(level, 9)


def get_toc_level(text: str) -> int:
    """Определение уровня оглавления на основе нумерации текста заголовка"""
    return get_heading_level_from_number(text)


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = BASE_DIR




def create_document(md_text: str) -> Document:
    """Создание документа из Markdown"""

    # Извлечение метаданных
    metadata = extract_metadata(md_text)

    # Парсинг ПЕРЕД очисткой (чтобы сохранить разделы)
    # Очищаем только служебные метаданные, но не разделы
    md_text_for_parsing = md_text
    md_text_for_parsing = re.sub(r'__?ФИО__?[:\-]?.*\n?', '', md_text_for_parsing)
    md_text_for_parsing = re.sub(r'__?Группа__?[:\-]?.*\n?', '', md_text_for_parsing)
    # Удаляем только служебные комментарии (с =), но НЕ разделы (без =)
    md_text_for_parsing = re.sub(r'\[//\]:\s*#\s*\w+\s*=.*\n?', '', md_text_for_parsing, flags=re.I)

    # Создание документа
    doc = Document()

    # Настройка страницы
    for s in doc.sections:
        s.page_height = Cm(DocumentSettings.PAGE_HEIGHT_CM)
        s.page_width = Cm(DocumentSettings.PAGE_WIDTH_CM)
        s.left_margin = Cm(DocumentSettings.LEFT_MARGIN_CM)
        s.right_margin = Cm(DocumentSettings.RIGHT_MARGIN_CM)
        s.top_margin = Cm(DocumentSettings.TOP_MARGIN_CM)
        s.bottom_margin = Cm(DocumentSettings.BOTTOM_MARGIN_CM)

    # Стиль Normal
    normal = doc.styles['Normal']
    normal.font.name = DocumentSettings.FONT_NAME
    normal.font.size = Pt(DocumentSettings.FONT_SIZE_PT)
    normal.paragraph_format.first_line_indent = Cm(DocumentSettings.FIRST_LINE_INDENT_CM)
    normal.paragraph_format.line_spacing = DocumentSettings.LINE_SPACING
    reset_heading_styles(doc)

    # Блок ФИО (если не таблица)
    if not metadata['is_table'] and (metadata['fio'] or metadata['group']):
        group_full = f"090302-{metadata['group']}" if metadata['group'] else ""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run1 = p.add_run(f"Выполнил студент: {metadata['fio']}\n")
        set_run_font(run1, size_pt=DocumentSettings.FONT_SIZE_PT, bold=True)
        run2 = p.add_run(f"Группа: {group_full}")
        set_run_font(run2, size_pt=DocumentSettings.FONT_SIZE_PT, bold=True)

    # Парсинг и сбор заголовков для оглавления
    parser = MarkdownParser(md_text_for_parsing)
    blocks = parser.parse()

    toc_entries = []

    def collect_headings(block):
        if isinstance(block, Section):
            for b in block.blocks:
                collect_headings(b)
        elif isinstance(block, HeadingBlock):
            toc_level = get_toc_level(block.text)
            toc_entries.append((toc_level, block.text))

    for block in blocks:
        collect_headings(block)

    # Рендеринг
    renderer = DocumentRenderer(doc, toc_entries, use_headings=metadata['use_headings'])
    for block in blocks:
        renderer.render_block(block)

    # Таблица подписей (если is_table)
    if metadata['is_table']:
        doc.add_page_break()

        # Вычисление места для таблицы
        sec = doc.sections[-1]
        page_height_cm = float(sec.page_height) / 360000.0
        top_cm = float(sec.top_margin) / 360000.0
        bottom_cm = float(sec.bottom_margin) / 360000.0
        usable_height_cm = page_height_cm - top_cm - bottom_cm

        total_cm = (
                           DocumentSettings.PAGE_WIDTH_CM - DocumentSettings.LEFT_MARGIN_CM - DocumentSettings.RIGHT_MARGIN_CM) * 0.75
        col_w_cm = total_cm / 2.0

        # Простая таблица подписей
        table = doc.add_table(rows=2, cols=2)
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for col in table.columns:
            for cell in col.cells:
                cell.width = Cm(col_w_cm)

        table.cell(0, 0).text = 'Выполнил:'
        table.cell(0, 1).text = metadata['fio']
        table.cell(1, 0).text = 'Проверил:'
        table.cell(1, 1).text = metadata['teacher']

        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:
                set_repeat_table_header(row)
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.style = doc.styles['Normal']
                    p.paragraph_format.first_line_indent = None
                    p.paragraph_format.left_indent = Cm(0)
                    p.paragraph_format.line_spacing = 1.5
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        set_run_font(run, size_pt=12)

        set_table_borders(table)

    # Номера страниц
    # Номера страниц
    if metadata['numerate']:
        add_page_number_centered(doc)

    return doc


def process_md_file(input_path: str):
    """Обработка одного файла"""
    filename = os.path.basename(input_path)
    name_wo_ext = os.path.splitext(filename)[0]
    output_path = os.path.join(BASE_DIR, f"{name_wo_ext}.docx")

    with open(input_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    doc = create_document(md_text)
    doc.save(output_path)
    print(f"✅ {filename} → {output_path}")


if __name__ == "__main__":
    # Поиск всех MD файлов
    all_md = [
        f for f in os.listdir(BASE_DIR)
        if f.lower().endswith(".md")
    ]

    if not all_md:
        print("⚠️ Не найдено ни одного Markdown-файла.")
    else:
        print(f"🔍 Найдено файлов: {len(all_md)}")
        for file in all_md:
            process_md_file(os.path.join(BASE_DIR, file))
        print("🎉 Все файлы обработаны.")
