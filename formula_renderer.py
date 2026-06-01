# -*- coding: utf-8 -*-
"""
Модуль рендеринга формул LaTeX -> OMML (нативный Word)

Принцип работы:
  1. Pandoc конвертирует LaTeX -> DOCX с нативным OMML
  2. Мы извлекаем XML-узел <m:oMathPara> из этого DOCX
  3. Вставляем его напрямую в параграф целевого документа

Результат: формулы редактируемы в редакторе формул Word,
           корректно отображаются без шрифтов/картинок.

ДИАГНОСТИКА: запустите этот файл напрямую для проверки окружения:
  python formula_renderer.py
"""

import os
import subprocess
import sys
import tempfile
import zipfile
from copy import deepcopy
from typing import Optional, List

from lxml import etree
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def add_paragraph_with_inline_formulas(doc: Document, text: str,
                                        align=None,
                                        first_line_indent_cm: float = 0,
                                        font_name: str = "Times New Roman",
                                        font_size_pt: int = 14) -> None:
    """
    Добавляет параграф, рендеря $...$ как OMML, остальное — как текст.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if first_line_indent_cm:
        p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)

    parts = re.split(r'\$(.+?)\$', text)

    for i, part in enumerate(parts):
        if i % 2 == 0:
            if part:
                run = p.add_run(part)
                run.font.name = font_name
                run.font.size = Pt(font_size_pt)
        else:
            # Пробуем OMML, иначе текстовый fallback
            omml_nodes = extract_omml(part)
            if omml_nodes:
                for node in omml_nodes:
                    p._p.append(deepcopy(node))
            else:
                run = p.add_run(f"${part}$")
                run.font.name = font_name
                run.font.size = Pt(font_size_pt)



# XML namespace для OMML
_OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# Кэш: None = не проверяли, True/False = результат проверки
_pandoc_available: Optional[bool] = None

import shutil


def _check_pandoc() -> bool:
    global _pandoc_available
    if _pandoc_available is None:
        # Ищем pandoc: сначала в PATH, потом в типичных местах установки
        pandoc_cmd = shutil.which("pandoc")

        if pandoc_cmd is None:
            # Типичные пути на Windows
            candidates = [
                r"C:\Program Files\Pandoc\pandoc.exe",
                os.path.expanduser(r"~\AppData\Local\Pandoc\pandoc.exe"),
            ]
            for path in candidates:
                if os.path.isfile(path):
                    pandoc_cmd = path
                    break

        if pandoc_cmd is None:
            _pandoc_available = False
            return False

        try:
            result = subprocess.run(
                [pandoc_cmd, "--version"],
                capture_output=True, timeout=5
            )
            _pandoc_available = result.returncode == 0
            # Сохраняем путь для дальнейших вызовов
            if _pandoc_available:
                global _pandoc_cmd
                _pandoc_cmd = pandoc_cmd
        except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
            _pandoc_available = False

    return _pandoc_available


# ---------------------------------------------------------------------------
# Публичный API
# ---------------------------------------------------------------------------

def extract_omml(latex: str) -> Optional[List]:
    """
    Конвертирует LaTeX-формулу в список OMML XML-элементов через pandoc.

    Возвращает список элементов для вставки в p._p,
    либо None если pandoc недоступен или произошла ошибка.
    """
    if not _check_pandoc():
        return None
    try:
        return _pandoc_latex_to_omml(latex)
    except Exception as e:
        # Логируем в stderr чтобы не ломать вывод
        print(f"[formula_renderer] extract_omml error: {e}", file=sys.stderr)
        return None


def add_formula_paragraph(doc: Document, latex: str,
                           number: Optional[str] = None,
                           font_size_pt: int = 14,
                           indent_cm: float = 3.0) -> bool:
    """
    Добавляет формулу в документ в формате ГОСТ:
      <формула по центру>    (номер)

    Если pandoc недоступен — добавляет текстовый fallback $latex$.

    Возвращает True если формула вставлена как OMML, False если fallback.
    """
    omml_nodes = extract_omml(latex)

    if omml_nodes:
        _insert_omml_table(doc, omml_nodes, number, font_size_pt)
        return True
    else:
        _insert_formula_fallback(doc, latex, number, font_size_pt)
        return False


def add_inline_formula(paragraph, latex: str) -> bool:
    """
    Вставляет inline-формулу в существующий параграф.

    Возвращает True если успешно, False если fallback (текст).
    """
    omml_nodes = extract_omml(latex)
    if omml_nodes:
        for node in omml_nodes:
            paragraph._p.append(deepcopy(node))
        return True
    else:
        run = paragraph.add_run(f"${latex}$")
        run.font.size = Pt(12)
        return False


# ---------------------------------------------------------------------------
# Внутренние функции
# ---------------------------------------------------------------------------

def _pandoc_latex_to_omml(latex: str) -> Optional[List]:
    pandoc_cmd = globals().get("_pandoc_cmd", "pandoc")  # ← добавить эту строку
    with tempfile.TemporaryDirectory() as tmpdir:
        md_path = os.path.join(tmpdir, "formula.md")
        docx_path = os.path.join(tmpdir, "formula.docx")

        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"$$\n{latex}\n$$\n")

        result = subprocess.run(
            [pandoc_cmd, md_path, "-o", docx_path],  # ← было ["pandoc", ...]
            capture_output=True,
            timeout=15
        )

        if result.returncode != 0 or not os.path.exists(docx_path):
            return None

        with zipfile.ZipFile(docx_path) as z:
            xml_bytes = z.read("word/document.xml")

    root = etree.fromstring(xml_bytes)

    # Ищем oMathPara — блочная формула (предпочтительно)
    nodes = root.findall(f".//{{{_OMML_NS}}}oMathPara")
    if nodes:
        return [deepcopy(nodes[0])]

    # Fallback: просто oMath (inline внутри pandoc)
    nodes = root.findall(f".//{{{_OMML_NS}}}oMath")
    if nodes:
        return [deepcopy(n) for n in nodes]

    return None


def _insert_omml_table(doc: Document, omml_nodes: List,
                        number: Optional[str],
                        font_size_pt: int):
    """
    Вставляет формулу в таблицу без границ:
      | <OMML по центру> | (N) по правому краю |
    """
    sec = doc.sections[0]
    usable_w = sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # Ширины: 85% формула, 15% номер
    w_formula = Cm(usable_w * 0.85)
    w_number = Cm(usable_w * 0.15)
    table.columns[0].width = w_formula
    table.columns[1].width = w_number

    # --- Левая ячейка: формула ---
    left = table.cell(0, 0)
    left_p = left.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_p.paragraph_format.space_before = Pt(6)
    left_p.paragraph_format.space_after = Pt(6)

    for node in omml_nodes:
        left_p._p.append(deepcopy(node))

    # --- Правая ячейка: номер ---
    right = table.cell(0, 1)
    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.paragraph_format.space_before = Pt(6)
    right_p.paragraph_format.space_after = Pt(6)

    if number:
        run = right_p.add_run(f"({number})")
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size_pt)

    # --- Убираем все границы ---
    _remove_table_borders(table)


def _remove_table_borders(table):
    """Убирает все границы таблицы через XML."""
    tbl = table._tbl
    tblPr = tbl.tblPr

    # Удаляем старые границы
    for child in list(tblPr):
        if child.tag.split("}")[-1] == "tblBorders":
            tblPr.remove(child)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _insert_formula_fallback(doc: Document, latex: str,
                              number: Optional[str],
                              font_size_pt: int):
    """Текстовый fallback если pandoc недоступен."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text = f"${latex}$"
    if number:
        text += f"    ({number})"
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size_pt)


# ---------------------------------------------------------------------------
# Диагностика при запуске напрямую
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    available = _check_pandoc()
    print(f"pandoc available: {'✅ YES' if available else '❌ NO'}")
    if available:
        pandoc_cmd = globals().get("_pandoc_cmd", "pandoc")  # ← добавить
        r = subprocess.run([pandoc_cmd, "--version"], capture_output=True)  # ← было ["pandoc", ...]
        print(f"  version: {r.stdout.decode().splitlines()[0]}")
    try:
        import lxml
        print(f"lxml available: ✅ YES (v{lxml.__version__})")
    except ImportError:
        print("lxml available: ❌ NO — run: pip install lxml")

    # 3. Тест формул
    test_cases = [
        ("Simple",    r"F = m \cdot a"),
        ("Fraction",  r"\frac{-b \pm \sqrt{b^2-4ac}}{2a}"),
        ("Integral",  r"\int_{-\infty}^{+\infty} e^{-x^2} dx = \sqrt{\pi}"),
        ("Matrix",    r"\begin{pmatrix} a & b \\ c & d \end{pmatrix}"),
        ("Greek",     r"-\frac{\hbar^2}{2m}\nabla^2\psi = E\psi"),
    ]

    print(f"\nFormula OMML extraction tests:")
    all_ok = True
    for name, latex in test_cases:
        nodes = extract_omml(latex)
        ok = nodes is not None and len(nodes) > 0
        all_ok = all_ok and ok
        print(f"  {name:12s}: {'✅ OMML' if ok else '❌ fallback'}")

    # 4. Генерация тестового документа
    if all_ok:
        print("\nGenerating test document...")
        doc = Document()
        for s in doc.sections:
            s.left_margin = Cm(3); s.right_margin = Cm(1)
            s.top_margin = Cm(2);  s.bottom_margin = Cm(2)

        for i, (name, latex) in enumerate(test_cases, 1):
            p = doc.add_paragraph(f"{name}: ")
            add_formula_paragraph(doc, latex, number=str(i))

        out = "formula_test_output.docx"
        doc.save(out)
        print(f"  ✅ Saved: {os.path.abspath(out)}")
    else:
        print("\n⚠️  Some formulas failed — check pandoc installation.")

    print("\n=== Done ===")