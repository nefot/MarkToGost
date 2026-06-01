# tests/test_formatting.py

import pytest
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from MarkToGost.utils.formatting import apply_italic_formatting, set_run_font, set_paragraph_formatting


# ================================
# apply_italic_formatting
# ================================

class TestApplyItalicFormatting:

    def test_plain_text(self):
        result = apply_italic_formatting("обычный текст")
        assert result == [("обычный текст", False)]

    def test_empty_string(self):
        result = apply_italic_formatting("")
        assert result == [("", False)]

    def test_none(self):
        result = apply_italic_formatting(None)
        assert result == [(None, False)]

    def test_only_italic(self):
        result = apply_italic_formatting("_курсив_")
        assert result == [("курсив", True)]

    def test_italic_in_middle(self):
        result = apply_italic_formatting("до _курсив_ после")
        assert result == [("до ", False), ("курсив", True), (" после", False)]

    def test_italic_at_start(self):
        result = apply_italic_formatting("_курсив_ после")
        assert result == [("курсив", True), (" после", False)]

    def test_italic_at_end(self):
        result = apply_italic_formatting("до _курсив_")
        assert result == [("до ", False), ("курсив", True)]

    def test_multiple_italic(self):
        result = apply_italic_formatting("_один_ и _два_")
        assert result == [("один", True), (" и ", False), ("два", True)]

    def test_no_underscore(self):
        result = apply_italic_formatting("текст без подчёркивания")
        assert result == [("текст без подчёркивания", False)]

    def test_single_underscore_ignored(self):
        # Одиночный _ без пары — не курсив
        result = apply_italic_formatting("текст_ без пары")
        assert result == [("текст_ без пары", False)]


# ================================
# set_run_font
# ================================

class TestSetRunFont:

    def _make_run(self):
        doc = Document()
        p = doc.add_paragraph()
        return p.add_run("тест")

    def test_default_font_name(self):
        run = self._make_run()
        set_run_font(run)
        assert run.font.name == "Times New Roman"

    def test_default_size(self):
        run = self._make_run()
        set_run_font(run)
        assert run.font.size == Pt(14)

    def test_custom_size(self):
        run = self._make_run()
        set_run_font(run, size_pt=12)
        assert run.font.size == Pt(12)

    def test_bold(self):
        run = self._make_run()
        set_run_font(run, bold=True)
        assert run.font.bold is True

    def test_italic(self):
        run = self._make_run()
        set_run_font(run, italic=True)
        assert run.font.italic is True

    def test_not_bold_by_default(self):
        run = self._make_run()
        set_run_font(run)
        assert run.font.bold is False

    def test_color_black(self):
        run = self._make_run()
        set_run_font(run)
        from docx.shared import RGBColor
        assert run.font.color.rgb == RGBColor(0, 0, 0)

