"""Рендеринг HTML-таблиц с поддержкой colspan/rowspan"""

from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from MarkToGost.config import DocumentSettings
from MarkToGost.formula_renderer import add_inline_formula
from MarkToGost.parser.blocks import HtmlTableBlock
from MarkToGost.utils.formatting import set_run_font
from MarkToGost.utils.xml_helpers import set_table_borders

from MarkToGost.parser.blocks import HtmlTableBlock, CellAlign
from docx.enum.text import WD_ALIGN_PARAGRAPH

_ALIGN_MAP = {
    CellAlign.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
    CellAlign.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
    CellAlign.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
}

def _remove_table_borders(table):
    """Делает таблицу прозрачной (без границ)"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    for child in list(tblPr):
        if child.tag.split("}")[-1] == "tblBorders":
            tblPr.remove(child)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)

def _merge_colspan(table, row_idx: int, col_idx: int, colspan: int):
    """Объединение ячеек по colspan"""
    if colspan <= 1:
        return
    start_cell = table.cell(row_idx, col_idx)
    end_cell = table.cell(row_idx, col_idx + colspan - 1)
    start_cell.merge(end_cell)


def render_html_table_block(renderer, block: HtmlTableBlock):
    """Рендеринг HTML-таблицы в DOCX"""
    if not block.rows:
        return

    # Определяем максимальное количество колонок
    max_cols = max(
        sum(cell.colspan for cell in row.cells)
        for row in block.rows
    )
    if max_cols <= 0:
        return

    sec = renderer.doc.sections[0]
    usable_width_cm = sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm
    col_width = Cm(usable_width_cm / max_cols)

    # Подпись
    if block.caption:
        cap = renderer.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.space_after = Pt(0)
        run = cap.add_run(f"Таблица {renderer.table_counter} — {block.caption}")
        set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, bold=False)

    # Создаём таблицу
    table = renderer.doc.add_table(rows=len(block.rows), cols=max_cols)
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Выставляем ширину колонок
    for col_idx in range(max_cols):
        for row in table.rows:
            row.cells[col_idx].width = col_width

    # Заполняем ячейки
    for r_idx, html_row in enumerate(block.rows):
        col_cursor = 0
        for cell_data in html_row.cells:
            # Пропускаем уже объединённые ячейки
            while col_cursor < max_cols and table.cell(r_idx, col_cursor).text != '':
                col_cursor += 1

            if col_cursor >= max_cols:
                break

            # Объединяем colspan
            if cell_data.colspan > 1:
                _merge_colspan(table, r_idx, col_cursor, cell_data.colspan)

            cell = table.cell(r_idx, col_cursor)
            cell.text = ""
            for p in cell.paragraphs:
                p.clear()

            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = DocumentSettings.LINE_SPACING
            p.alignment = _ALIGN_MAP.get(cell_data.align, WD_ALIGN_PARAGRAPH.LEFT)

            if cell_data.text:
                run = p.add_run(cell_data.text)
                set_run_font(
                    run,
                    size_pt=DocumentSettings.TABLE_FONT_SIZE_PT,
                    bold=cell_data.bold,
                    italic=cell_data.italic,
                )
                if cell_data.underline:
                    run.underline = True

            if cell_data.formula:
                add_inline_formula(p, cell_data.formula)

            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            col_cursor += cell_data.colspan

    if block.transparent:
        _remove_table_borders(table)
    else:
        set_table_borders(table)
    renderer.table_counter += 1
    renderer._mark_content()