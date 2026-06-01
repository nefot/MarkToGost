"""Рендеринг таблиц"""

from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from MarkToGost.config import DocumentSettings
from MarkToGost.parser.blocks import TableBlock
from MarkToGost.utils.formatting import apply_italic_formatting, set_run_font, set_paragraph_formatting
from MarkToGost.utils.xml_helpers import set_table_borders, set_repeat_table_header
from MarkToGost.utils.document_helpers import split_md_table_row


def _estimate_row_height_cm(row_data: list, col_width_cm: float, size_pt: int = 12) -> float:
    """Оценка высоты строки таблицы"""
    LINE_HEIGHT_CM = size_pt * 0.0353  # pt -> cm
    CELL_PADDING_CM = 0.15  # внутренние отступы ячейки
    SAFETY_FACTOR = 1.05  # минимальный запас

    # Кириллица шире латиницы — используем 0.025 вместо 0.022/0.030
    CHAR_WIDTH_CM = size_pt * 0.025

    max_lines = 1
    for cell_text in row_data:
        if not cell_text:
            continue
        chars_per_line = max(1, int(col_width_cm / CHAR_WIDTH_CM))
        lines = 0
        for paragraph in str(cell_text).split('\n'):
            lines += max(1, -(-len(paragraph) // chars_per_line))  # ceiling division
        max_lines = max(max_lines, lines)

    return max_lines * LINE_HEIGHT_CM * DocumentSettings.LINE_SPACING * SAFETY_FACTOR + CELL_PADDING_CM


def render_table_block(renderer, block: TableBlock):
    """Рендеринг таблицы"""
    if not block.rows:
        return

    parsed_rows = [split_md_table_row(row) for row in block.rows]
    max_cols = max(len(r) for r in parsed_rows) if parsed_rows else 0
    if max_cols <= 0:
        return

    for row in parsed_rows:
        if len(row) < max_cols:
            row.extend([""] * (max_cols - len(row)))

    # Вычисляем доступную высоту страницы
    sec = renderer.doc.sections[0]
    usable_width_cm = sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm
    PAGE_MARGIN_CM = 0.8  # только колонтитул с номером страницы
    usable_height_cm = sec.page_height.cm - sec.top_margin.cm - sec.bottom_margin.cm - PAGE_MARGIN_CM
    col_width_cm = usable_width_cm / max_cols

    CAPTION_HEIGHT_CM = 0.8  # высота строки подписи "Таблица N"
    HEADER_HEIGHT_CM = _estimate_row_height_cm(
        parsed_rows[0], col_width_cm, DocumentSettings.TABLE_FONT_SIZE_PT
    )

    # Разбиваем на chunks по реальной высоте
    chunks = []
    current_chunk = []
    # Первый chunk: вычитаем высоту подписи + заголовка
    current_height = CAPTION_HEIGHT_CM + HEADER_HEIGHT_CM
    header_row = parsed_rows[0]

    for row_data in parsed_rows:
        row_h = _estimate_row_height_cm(
            row_data, col_width_cm, DocumentSettings.TABLE_FONT_SIZE_PT
        )

        if current_chunk and current_height + row_h > usable_height_cm:
            chunks.append(current_chunk)
            current_chunk = [header_row]  # повторяем заголовок в каждом chunk
            # Следующий chunk: вычитаем высоту подписи "Продолжение" + заголовка
            current_height = CAPTION_HEIGHT_CM + HEADER_HEIGHT_CM + row_h
        else:
            current_height += row_h
            current_chunk.append(row_data)

    if current_chunk:
        chunks.append(current_chunk)

    col_width = Cm(col_width_cm)

    for chunk_index, chunk in enumerate(chunks):
        # --- Подпись ---
        cap = renderer.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_formatting(
            cap,
            first_line_indent=None,
            left_indent=Cm(0),
            space_before=6,
            space_after=6 if chunk_index == len(chunks) - 1 else 0,  # Интервал после последней таблицы
            line_spacing=DocumentSettings.LINE_SPACING
        )

        if chunk_index == 0:
            caption_full = (
                f"Таблица {renderer.table_counter} — {block.caption}"
                if block.caption
                else f"Таблица {renderer.table_counter}"
            )
        else:
            caption_full = f"Продолжение таблицы {renderer.table_counter}"

        cap.clear()
        for part_text, is_italic in apply_italic_formatting(caption_full):
            run = cap.add_run(part_text)
            set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, bold=False, italic=is_italic)

        # --- Таблица ---
        table = renderer.doc.add_table(rows=len(chunk), cols=max_cols)
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False

        for col_idx in range(max_cols):
            for row in table.rows:
                row.cells[col_idx].width = col_width

        # Повтор заголовка средствами Word (на случай если наша оценка немного ошиблась)
        if chunk:
            set_repeat_table_header(table.rows[0])

        # --- Заполнение ячеек ---
        for r_idx, row_data in enumerate(chunk):
            row = table.rows[r_idx]

            # Запрет разрыва строки посередине
            trPr = row._tr.get_or_add_trPr()
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '1')
            trPr.append(cantSplit)

            for c_idx, value in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = ""
                for p in cell.paragraphs:
                    p.clear()

                p = cell.paragraphs[0]
                p.style = None
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.left_indent = Pt(0)
                p.paragraph_format.right_indent = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = DocumentSettings.LINE_SPACING
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT

                for part_text, is_italic in apply_italic_formatting(value if value is not None else ""):
                    run = p.add_run(part_text)
                    set_run_font(
                        run,
                        size_pt=DocumentSettings.TABLE_FONT_SIZE_PT,
                        bold=r_idx == 0,
                        italic=is_italic
                    )

                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        set_table_borders(table)

        if chunk_index < len(chunks) - 1:
            renderer.doc.add_page_break()

    renderer.table_counter += 1
    renderer._mark_content()

