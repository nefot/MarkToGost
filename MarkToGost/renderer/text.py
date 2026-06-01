"""Рендеринг текстовых блоков и заголовков"""

import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from MarkToGost.config import DocumentSettings
from MarkToGost.parser.blocks import TextBlock, HeadingBlock
from MarkToGost.utils.formatting import apply_italic_formatting, set_run_font, set_paragraph_formatting
from MarkToGost.utils.document_helpers import replace_image_refs
from docx.shared import Cm


def render_text_block(renderer, block: TextBlock):
    """Рендеринг текстового блока с заменой ссылок на изображения"""
    # Заменяем ссылки @img_id на номера
    text = replace_image_refs(block.text, renderer.image_refs)

    p = renderer.doc.add_paragraph(text)
    set_paragraph_formatting(
        p,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(DocumentSettings.FIRST_LINE_INDENT_CM),
        line_spacing=DocumentSettings.LINE_SPACING
    )

    # Применяем курсив
    p.clear()
    for part_text, is_italic in apply_italic_formatting(text):
        run = p.add_run(part_text)
        set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, italic=is_italic)

    renderer._mark_content()


def render_heading_block(renderer, block: HeadingBlock):
    """Рендеринг заголовка"""
    if not renderer.use_headings:
        # Рендерим как обычный жирный текст без стиля Heading
        p = renderer.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if block.level > 1 else WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_formatting(p, space_before=12, space_after=6, line_spacing=DocumentSettings.LINE_SPACING)
        text = block.text.upper() if block.level == 1 else block.text
        for part_text, is_italic in apply_italic_formatting(text):
            run = p.add_run(part_text)
            set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, bold=True, italic=is_italic)
        renderer._mark_content()
        return

    level = min(block.level, 9)
    style_name = f'Heading {level}'
    p = renderer.doc.add_paragraph(style=style_name)

    if level == 1:
        set_paragraph_formatting(p, space_before=0, space_after=6)
    else:
        set_paragraph_formatting(p, space_before=12, space_after=6)

    p.paragraph_format.outline_level = level - 1

    # Принудительно задаём выравнивание через XML
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:jc')):
        pPr.remove(old)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center' if level == 1 else 'left')
    pPr.append(jc)

    # Текст: uppercase только для H1
    text = block.text.upper() if level == 1 else block.text

    p.clear()
    for part_text, is_italic in apply_italic_formatting(text):
        run = p.add_run(part_text)
        set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, bold=True, italic=is_italic)
        # Принудительно отключаем caps через XML
        rPr = run._r.get_or_add_rPr()
        for old in rPr.findall(qn('w:caps')):
            rPr.remove(old)
        for old in rPr.findall(qn('w:smallCaps')):
            rPr.remove(old)
        caps = OxmlElement('w:caps')
        caps.set(qn('w:val'), '0')
        rPr.append(caps)

    renderer._mark_content()

