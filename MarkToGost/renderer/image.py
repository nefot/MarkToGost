"""Рендеринг блоков изображений"""

import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

from MarkToGost.config import DocumentSettings
from MarkToGost.parser.blocks import ImageBlock
from MarkToGost.utils.document_helpers import compute_image_width_cm
from MarkToGost.utils.formatting import apply_italic_formatting, set_run_font, set_paragraph_formatting

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def render_image_block(renderer, block: ImageBlock):
    """Рендеринг изображения"""
    img_path = block.path
    if not os.path.isabs(img_path):
        img_path = os.path.join(BASE_DIR, block.path)

    if os.path.exists(img_path):
        # Изображение
        pic_para = renderer.doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        img_width_cm = compute_image_width_cm(renderer.doc, DocumentSettings.IMAGE_WIDTH_FRACTION)
        try:
            pic_para.add_run().add_picture(img_path, width=img_width_cm)
        except Exception:
            pic_para.add_run(f"[Ошибка вставки: {os.path.basename(img_path)}]")

        # Получаем номер изображения
        # Если img_id указан и был собран на первом проходе, используем сохраненный номер
        if block.img_id and block.img_id in renderer.image_refs:
            fig_num = renderer.image_refs[block.img_id]
        else:
            # Иначе увеличиваем счетчик
            renderer.figure_counter += 1
            fig_num = renderer.figure_counter
            # Сохраняем в image_refs, если есть ID
            if block.img_id:
                renderer.image_refs[block.img_id] = fig_num

        fname = os.path.basename(block.path)
        caption_text = block.caption or ""

        if caption_text:
            caption_full = f"Рисунок {fig_num} - {caption_text.strip().capitalize()}"
        else:
            caption_full = f"Рисунок {fig_num}"

        cap_para = renderer.doc.add_paragraph(caption_full)
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_formatting(cap_para, line_spacing=DocumentSettings.LINE_SPACING)

        # Применяем курсив
        cap_para.clear()
        for part_text, is_italic in apply_italic_formatting(caption_full):
            run = cap_para.add_run(part_text)
            set_run_font(run, size_pt=DocumentSettings.CAPTION_FONT_SIZE_PT,
                        italic=DocumentSettings.CAPTION_ITALIC or is_italic)

        renderer.image_map[fname] = str(fig_num)
        renderer._mark_content()
    else:
        p = renderer.doc.add_paragraph(f"[Изображение не найдено: {os.path.basename(img_path)}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        renderer._mark_content()

