"""Рендеринг списков (упорядоченных и неупорядоченных)"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from MarkToGost.config import DocumentSettings
from MarkToGost.parser.blocks import ListBlock
from MarkToGost.utils.formatting import apply_italic_formatting, set_run_font, set_paragraph_formatting


def render_list_block(renderer, block: ListBlock):
    """Рендеринг списка"""
    for i, item in enumerate(block.items, 1):
        if block.ordered:
            text = f"{i}. {item}"
        else:
            text = f"– {item}"

        p = renderer.doc.add_paragraph()

        if block.ordered:
            set_paragraph_formatting(
                p,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=Cm(DocumentSettings.FIRST_LINE_INDENT_CM),
                line_spacing=DocumentSettings.LINE_SPACING
            )

            # Номер
            run = p.add_run(f"{i}. ")
            set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT)

            # Текст с курсивом
            for part_text, is_italic in apply_italic_formatting(item):
                run = p.add_run(part_text)
                set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, italic=is_italic)

        else:
            set_paragraph_formatting(
                p,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                left_indent=Cm(DocumentSettings.FIRST_LINE_INDENT_CM * 2),
                first_line_indent=Cm(-0.35),
                line_spacing=DocumentSettings.LINE_SPACING,
                space_before=0,
                space_after=0 if i < len(block.items) else 6
            )

            # Маркер
            run = p.add_run("– ")
            set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT)

            # Текст с курсивом
            for part_text, is_italic in apply_italic_formatting(item):
                run = p.add_run(part_text)
                set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, italic=is_italic)

    renderer._mark_content()

