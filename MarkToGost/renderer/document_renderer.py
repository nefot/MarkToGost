from typing import List
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

from MarkToGost.config import DocumentSettings
from MarkToGost.formula_renderer import add_paragraph_with_inline_formulas, add_formula_paragraph
from MarkToGost.parser.blocks import (
    BaseBlock, Section, TextBlock, HeadingBlock, ImageBlock, ListBlock,
    TableBlock, CodeBlock, FormulaBlock
)
from MarkToGost.utils.formatting import apply_italic_formatting, set_run_font, set_paragraph_formatting
from MarkToGost.utils.xml_helpers import set_table_borders, set_repeat_table_header
from MarkToGost.utils.toc import add_toc, get_heading_level_from_number
from MarkToGost.utils.document_helpers import (
    compute_image_width_cm, replace_image_refs, split_md_table_row
)

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


class DocumentRenderer:
    """Рендерер блоков в DOCX"""

    def __init__(self, doc: Document, toc_entries: List[tuple] = None, use_headings: bool = True):
        self.doc = doc
        self.table_counter = 1
        self.formula_counter = 1  # Счетчик формул
        self.figure_counter = 0
        self.image_map = {}
        self.image_refs = {}  # Для хранения ссылок на изображения по их ID
        self.formula_refs = {}  # Для хранения ссылок на формулы по их ID
        self.first_h1_seen = False
        self.first_h4_seen = False
        self.chapter_num = 0
        self.section_num = 0
        self.use_headings = use_headings
        self.subsection_num = 0
        self.toc_entries = toc_entries or []
        self._last_was_break = False  # Флаг для защиты от двойных разрывов

    def _is_document_start(self) -> bool:
        """Проверка начала документа"""
        return len(self.doc.paragraphs) == 0

    def _safe_page_break(self):
        """Не допускает двойных пустых страниц"""
        if not self._last_was_break:
            self.doc.add_page_break()
            self._last_was_break = True

    def _mark_content(self):
        """Сброс флага page break после добавления контента"""
        self._last_was_break = False

    def render_block(self, block: BaseBlock):
        """Рендеринг блока"""
        if isinstance(block, Section):
            self._render_section_block(block)
        elif isinstance(block, TextBlock):
            self._render_text_block(block)
        elif isinstance(block, HeadingBlock):
            self._render_heading_block(block)
        elif isinstance(block, ImageBlock):
            self._render_image_block(block)
        elif isinstance(block, ListBlock):
            self._render_list_block(block)
        elif isinstance(block, TableBlock):
            self._render_table_block(block)
        elif isinstance(block, CodeBlock):
            self._render_code_block(block)
        elif isinstance(block, FormulaBlock):
            self._render_formula_block(block)

    def _render_section_block(self, section: Section):
        """Рендеринг раздела с его блоками
        Строгое правило: раздел ВСЕГДА начинается и заканчивается разрывом (если включено)
        """
        if not section.add_page_breaks:
            # Если page breaks отключены, просто рендерим содержимое
            if section.section_id:
                if section.section_id == "ОГЛАВЛЕНИЕ":
                    # Обычный текст, НЕ Heading (чтобы не попало в TOC)
                    p = self.doc.add_paragraph()
                    # Явный сброс форматирования стиля
                    p.style.font.all_caps = False
                    p.style.font.bold = False
                    p.paragraph_format.alignment = None  # сбросим до явной установки ниже

                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    run = p.add_run(section.section_id)
                    set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, bold=True)

                    self._mark_content()
                else:
                    level = get_heading_level_from_number(section.section_id)
                    heading_block = HeadingBlock(text=section.section_id, level=level)
                    self._render_heading_block(heading_block)

            for block in section.blocks:
                self.render_block(block)
            return

        # 🔹 ВСЕГДА начинаем с новой страницы (кроме самого начала документа)
        if not self._is_document_start():
            self._safe_page_break()

        # 🔹 Заголовок раздела
        if section.section_id:
            if section.section_id == "ОГЛАВЛЕНИЕ":
                # Обычный текст, НЕ Heading (чтобы не попало в TOC)
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_formatting(p, space_before=0, space_after=6)

                run = p.add_run(section.section_id)
                set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, bold=True)

                self._mark_content()
            else:
                level = get_heading_level_from_number(section.section_id)
                heading_block = HeadingBlock(text=section.section_id, level=level)
                self._render_heading_block(heading_block)

        # 🔹 Специальная обработка для оглавления
        if section.section_id == "ОГЛАВЛЕНИЕ":
            add_toc(self.doc)
            self._mark_content()

        # 🔹 Контент раздела
        for block in section.blocks:
            self.render_block(block)

        # 🔹 ВСЕГДА завершаем раздел разрывом страницы
        self._safe_page_break()

    def _render_formula_block(self, block):
        """
        Рендеринг блочной формулы с номером по ГОСТ.

        Использует нативный OMML через pandoc — формулы редактируемы
        в редакторе формул Word, не требуют внешних шрифтов.
        """
        import sys
        import os
        # Добавляем родительскую папку в path для импорта formula_renderer
        parent_dir = os.path.dirname(os.path.dirname(BASE_DIR))
        if parent_dir not in sys.path:
            sys.path.insert(0, parent_dir)
        

        # --- Нумерация ---
        if not block.number:
            block.number = str(self.formula_counter)
            self.formula_counter += 1

        if block.formula_id:
            self.formula_refs[block.formula_id] = block.number

        # --- Вставка формулы ---
        add_formula_paragraph(
            self.doc,
            latex=block.latex,
            number=block.number,
            font_size_pt=DocumentSettings.FONT_SIZE_PT,
        )

        # --- Пояснение (где ...) ---
        if block.explanation:
            for raw_line in block.explanation.strip().split('\\n'):
                raw_line = raw_line.strip()
                if not raw_line:
                    continue

                add_paragraph_with_inline_formulas(
                    self.doc,
                    text=raw_line,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    first_line_indent_cm=DocumentSettings.FIRST_LINE_INDENT_CM,
                    font_size_pt=DocumentSettings.FONT_SIZE_PT,
                )

        self.doc.add_paragraph()  # отступ после формулы
        self._mark_content()

    def _render_text_block(self, block: TextBlock):
        """Рендеринг текстового блока с заменой ссылок на изображения"""
        # Заменяем ссылки @img_id на номера
        text = replace_image_refs(block.text, self.image_refs)

        p = self.doc.add_paragraph(text)
        set_paragraph_formatting(
            p,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=Cm(DocumentSettings.FIRST_LINE_INDENT_CM),
            line_spacing=DocumentSettings.LINE_SPACING
        )

        # Применяем курсив
        p.clear()
        for part_text, is_italic in apply_italic_formatting(text):
            run = p.add_run(part_text)
            set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, italic=is_italic)

        self._mark_content()

    def _render_heading_block(self, block: HeadingBlock):
        if not self.use_headings:
            # Рендерим как обычный жирный текст без стиля Heading
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if block.level > 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_formatting(p, space_before=12, space_after=6, line_spacing=DocumentSettings.LINE_SPACING)
            text = block.text.upper() if block.level == 1 else block.text
            for part_text, is_italic in apply_italic_formatting(text):
                run = p.add_run(part_text)
                set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, bold=True, italic=is_italic)
            self._mark_content()
            return
        # ... остальной код метода без изменений
        level = min(block.level, 9)
        style_name = f'Heading {level}'
        p = self.doc.add_paragraph(style=style_name)

        if level == 1:
            set_paragraph_formatting(p, space_before=0, space_after=6)
        else:
            set_paragraph_formatting(p, space_before=12, space_after=6)

        p.paragraph_format.outline_level = level - 1

        # Принудительно задаём выравнивание через XML
        pPr = p._p.get_or_add_pPr()
        for old in pPr.findall(qn('w:jc')):
            pPr.remove(old)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center' if level == 1 else 'left')
        pPr.append(jc)

        # Текст: uppercase только для H1
        text = block.text.upper() if level == 1 else block.text

        p.clear()
        for part_text, is_italic in apply_italic_formatting(text):
            run = p.add_run(part_text)
            set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, bold=True, italic=is_italic)
            # Принудительно отключаем caps через XML
            rPr = run._r.get_or_add_rPr()
            for old in rPr.findall(qn('w:caps')):
                rPr.remove(old)
            for old in rPr.findall(qn('w:smallCaps')):
                rPr.remove(old)
            caps = OxmlElement('w:caps')
            caps.set(qn('w:val'), '0')
            rPr.append(caps)

        self._mark_content()

    def _render_image_block(self, block: ImageBlock):
        """Рендеринг изображения"""
        img_path = block.path
        if not os.path.isabs(img_path):
            img_path = os.path.join(BASE_DIR, block.path)

        if os.path.exists(img_path):
            # Изображение
            pic_para = self.doc.add_paragraph()
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            img_width_cm = compute_image_width_cm(self.doc, DocumentSettings.IMAGE_WIDTH_FRACTION)
            try:
                pic_para.add_run().add_picture(img_path, width=img_width_cm)
            except Exception:
                pic_para.add_run(f"[Ошибка вставки: {os.path.basename(img_path)}]")

            # Получаем номер изображения
            # Если img_id указан и был собран на первом проходе, используем сохраненный номер
            if block.img_id and block.img_id in self.image_refs:
                fig_num = self.image_refs[block.img_id]
            else:
                # Иначе увеличиваем счетчик
                self.figure_counter += 1
                fig_num = self.figure_counter
                # Сохраняем в image_refs, если есть ID
                if block.img_id:
                    self.image_refs[block.img_id] = fig_num

            fname = os.path.basename(block.path)
            caption_text = block.caption or ""

            if caption_text:
                caption_full = f"Рисунок {fig_num} - {caption_text.strip().capitalize()}"
            else:
                caption_full = f"Рисунок {fig_num}"

            cap_para = self.doc.add_paragraph(caption_full)
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_formatting(cap_para, line_spacing=DocumentSettings.LINE_SPACING)

            # Применяем курсив
            cap_para.clear()
            for part_text, is_italic in apply_italic_formatting(caption_full):
                run = cap_para.add_run(part_text)
                set_run_font(run, size_pt=DocumentSettings.CAPTION_FONT_SIZE_PT,
                             italic=DocumentSettings.CAPTION_ITALIC or is_italic)

            self.image_map[fname] = str(fig_num)
            self._mark_content()
        else:
            p = self.doc.add_paragraph(f"[Изображение не найдено: {os.path.basename(img_path)}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._mark_content()

    def _render_list_block(self, block: ListBlock):
        """Рендеринг списка"""
        for i, item in enumerate(block.items, 1):
            if block.ordered:
                text = f"{i}. {item}"
            else:
                text = f"– {item}"

            p = self.doc.add_paragraph()

            if block.ordered:
                set_paragraph_formatting(
                    p,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line_indent=Cm(DocumentSettings.FIRST_LINE_INDENT_CM),
                    line_spacing=DocumentSettings.LINE_SPACING
                )

                # Номер
                run = p.add_run(f"{i}. ")
                set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT)

                # Текст с курсивом
                for part_text, is_italic in apply_italic_formatting(item):
                    run = p.add_run(part_text)
                    set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, italic=is_italic)

            else:
                set_paragraph_formatting(
                    p,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    left_indent=Cm(DocumentSettings.FIRST_LINE_INDENT_CM * 2),
                    first_line_indent=Cm(-0.35),
                    line_spacing=DocumentSettings.LINE_SPACING,
                    space_before=0,
                    space_after=0 if i < len(block.items) else 6
                )

                # Маркер
                run = p.add_run("– ")
                set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT)

                # Текст с курсивом
                for part_text, is_italic in apply_italic_formatting(item):
                    run = p.add_run(part_text)
                    set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, italic=is_italic)

        self._mark_content()

    def _render_code_block(self, block: CodeBlock):
        """Рендеринг блока кода"""
        code = block.code.strip()
        language = block.language.strip().lower()

        if language == "python":
            # Специфическое форматирование для Python
            code = re.sub(r'^\s*def\s+(\w+)\s*\(', r'    def \1(', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*class\s+(\w+)\s*\(', r'    class \1(', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*#', r'    #', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*print\s*\(', r'    print(', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*return\s+', r'        return ', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*if\s+', r'    if ', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*else\s+', r'    else ', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*elif\s+', r'    elif ', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*for\s+', r'    for ', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*while\s+', r'    while ', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*try\s+', r'    try:', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*except\s+', r'    except:', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*finally\s+', r'    finally:', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*with\s+', r'    with ', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*as\s+', r'    as ', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*await\s+', r'    await ', code, flags=re.MULTILINE)
            code = re.sub(r'^\s*async\s+', r'    async ', code, flags=re.MULTILINE)

        # Общие правила для всех языков
        code = re.sub(r'^\s*//\s*', r'    // ', code, flags=re.MULTILINE)
        code = re.sub(r'^\s*#\s*', r'    # ', code, flags=re.MULTILINE)
        code = re.sub(r'^\s*;\s*', r'    ;', code, flags=re.MULTILINE)
        code = re.sub(r'^\s*{\s*', r'    {', code, flags=re.MULTILINE)
        code = re.sub(r'^\s*}\s*', r'    }', code, flags=re.MULTILINE)
        code = re.sub(r'^\s*\(\s*', r'    (', code, flags=re.MULTILINE)
        code = re.sub(r'^\s*\)\s*', r'    )', code, flags=re.MULTILINE)

        # Удаление пустых строк
        code = re.sub(r'^\s*\n', '', code, flags=re.MULTILINE)

        # Добавление блока кода
        p = self.doc.add_paragraph()
        p.add_run(code).font.name = 'Courier New'
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        self._mark_content()

    def _estimate_row_height_cm(self, row_data: list, col_width_cm: float, size_pt: int = 12) -> float:
        LINE_HEIGHT_CM = size_pt * 0.0353  # pt -> cm
        CELL_PADDING_CM = 0.15  # внутренние отступы ячейки
        SAFETY_FACTOR = 1.05  # минимальный запас

        # Кириллица шире латиницы — используем 0.025 вместо 0.022/0.030
        CHAR_WIDTH_CM = size_pt * 0.025

        max_lines = 1
        for cell_text in row_data:
            if not cell_text:
                continue
            chars_per_line = max(1, int(col_width_cm / CHAR_WIDTH_CM))
            lines = 0
            for paragraph in str(cell_text).split('\n'):
                lines += max(1, -(-len(paragraph) // chars_per_line))  # ceiling division
            max_lines = max(max_lines, lines)

        return max_lines * LINE_HEIGHT_CM * DocumentSettings.LINE_SPACING * SAFETY_FACTOR + CELL_PADDING_CM

    def _render_table_block(self, block: TableBlock):
        if not block.rows:
            return

        parsed_rows = [split_md_table_row(row) for row in block.rows]
        max_cols = max(len(r) for r in parsed_rows) if parsed_rows else 0
        if max_cols <= 0:
            return

        for row in parsed_rows:
            if len(row) < max_cols:
                row.extend([""] * (max_cols - len(row)))

        # Вычисляем доступную высоту страницы
        sec = self.doc.sections[0]
        usable_width_cm = sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm
        PAGE_MARGIN_CM = 0.8  # только колонтитул с номером страницы
        usable_height_cm = sec.page_height.cm - sec.top_margin.cm - sec.bottom_margin.cm - PAGE_MARGIN_CM
        col_width_cm = usable_width_cm / max_cols

        CAPTION_HEIGHT_CM = 0.8  # высота строки подписи "Таблица N"
        HEADER_HEIGHT_CM = self._estimate_row_height_cm(
            parsed_rows[0], col_width_cm, DocumentSettings.TABLE_FONT_SIZE_PT
        )

        # Разбиваем на chunks по реальной высоте
        chunks = []
        current_chunk = []
        # Первый chunk: вычитаем высоту подписи + заголовка
        current_height = CAPTION_HEIGHT_CM + HEADER_HEIGHT_CM
        header_row = parsed_rows[0]

        for row_data in parsed_rows:
            row_h = self._estimate_row_height_cm(
                row_data, col_width_cm, DocumentSettings.TABLE_FONT_SIZE_PT
            )

            if current_chunk and current_height + row_h > usable_height_cm:
                chunks.append(current_chunk)
                current_chunk = [header_row]  # повторяем заголовок в каждом chunk
                # Следующий chunk: вычитаем высоту подписи "Продолжение" + заголовка
                current_height = CAPTION_HEIGHT_CM + HEADER_HEIGHT_CM + row_h
            else:
                current_height += row_h
                current_chunk.append(row_data)

        if current_chunk:
            chunks.append(current_chunk)

        col_width = Cm(col_width_cm)

        for chunk_index, chunk in enumerate(chunks):
            # --- Подпись ---
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_paragraph_formatting(
                cap,
                first_line_indent=None,
                left_indent=Cm(0),
                space_before=6,
                space_after=0,
                line_spacing=DocumentSettings.LINE_SPACING
            )

            if chunk_index == 0:
                caption_full = (
                    f"Таблица {self.table_counter} — {block.caption}"
                    if block.caption
                    else f"Таблица {self.table_counter}"
                )
            else:
                caption_full = f"Продолжение таблицы {self.table_counter}"

            cap.clear()
            for part_text, is_italic in apply_italic_formatting(caption_full):
                run = cap.add_run(part_text)
                set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, bold=False, italic=is_italic)

            # --- Таблица ---
            table = self.doc.add_table(rows=len(chunk), cols=max_cols)
            try:
                table.style = 'Table Grid'
            except Exception:
                pass
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False

            for col_idx in range(max_cols):
                for row in table.rows:
                    row.cells[col_idx].width = col_width

            # Повтор заголовка средствами Word (на случай если наша оценка немного ошиблась)
            if chunk:
                set_repeat_table_header(table.rows[0])

            # --- Заполнение ячеек ---
            for r_idx, row_data in enumerate(chunk):
                row = table.rows[r_idx]

                # Запрет разрыва строки посередине
                trPr = row._tr.get_or_add_trPr()
                cantSplit = OxmlElement('w:cantSplit')
                cantSplit.set(qn('w:val'), '1')
                trPr.append(cantSplit)

                for c_idx, value in enumerate(row_data):
                    cell = row.cells[c_idx]
                    cell.text = ""
                    for p in cell.paragraphs:
                        p.clear()

                    p = cell.paragraphs[0]
                    p.style = None
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.left_indent = Pt(0)
                    p.paragraph_format.right_indent = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = DocumentSettings.LINE_SPACING
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT

                    for part_text, is_italic in apply_italic_formatting(value if value is not None else ""):
                        run = p.add_run(part_text)
                        set_run_font(
                            run,
                            size_pt=DocumentSettings.TABLE_FONT_SIZE_PT,
                            bold=r_idx == 0,
                            italic=is_italic
                        )

                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            set_table_borders(table)

            if chunk_index < len(chunks) - 1:
                self.doc.add_page_break()

        self.table_counter += 1
        self._mark_content()
