from MarkToGost.parser.blocks import Section, HeadingBlock
from docx.enum.text import WD_ALIGN_PARAGRAPH

from MarkToGost.config import DocumentSettings
from MarkToGost.utils.formatting import set_run_font, set_paragraph_formatting
from MarkToGost.utils.toc import add_toc, get_heading_level_from_number
from MarkToGost.renderer.text import render_heading_block


def render_section_block(renderer, section: Section):
    """Рендеринг раздела с его блоками
    Строгое правило: раздел ВСЕГДА начинается и заканчивается разрывом (если включено)
    """
    if not section.add_page_breaks:
        # Если page breaks отключены, просто рендерим содержимое
        if section.section_id:
            if section.section_id == "ОГЛАВЛЕНИЕ":
                # Обычный текст, НЕ Heading (чтобы не попало в TOC)
                p = renderer.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(section.section_id)
                set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, bold=True)
                renderer._mark_content()
            else:
                level = get_heading_level_from_number(section.section_id)
                heading_block = HeadingBlock(text=section.section_id, level=level)
                render_heading_block(renderer, heading_block)

        for block in section.blocks:
            renderer.render_block(block)
        return

    # 🔹 ВСЕГДА начинаем с новой страницы (кроме самого начала документа)
    if not renderer._is_document_start():
        renderer._safe_page_break()

    # 🔹 Заголовок раздела
    if section.section_id:
        if section.section_id == "ОГЛАВЛЕНИЕ":
            # Обычный текст, НЕ Heading (чтобы не попало в TOC)
            p = renderer.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_formatting(p, space_before=0, space_after=6)
            run = p.add_run(section.section_id)
            set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT, bold=True)
            renderer._mark_content()
        else:
            level = get_heading_level_from_number(section.section_id)
            heading_block = HeadingBlock(text=section.section_id, level=level)
            render_heading_block(renderer, heading_block)

    # 🔹 Специальная обработка для оглавления
    if section.section_id == "ОГЛАВЛЕНИЕ":
        add_toc(renderer.doc)
        renderer._mark_content()

    # 🔹 Контент раздела
    for block in section.blocks:
        renderer.render_block(block)

    # 🔹 ВСЕГДА завершаем раздел разрывом страницы
    renderer._safe_page_break()
