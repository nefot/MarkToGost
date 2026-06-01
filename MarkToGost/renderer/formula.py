"""Рендеринг блоков формул"""

from MarkToGost.config import DocumentSettings
from MarkToGost.parser.blocks import FormulaBlock
from MarkToGost.formula_renderer import add_formula_paragraph, add_paragraph_with_inline_formulas
from docx.enum.text import WD_ALIGN_PARAGRAPH


def render_formula_block(renderer, block: FormulaBlock):
    """
    Рендеринг блочной формулы с номером по ГОСТ.

    Использует нативный OMML через pandoc — формулы редактируемы
    в редакторе формул Word, не требуют внешних шрифтов.
    """
    # --- Нумерация ---
    if not block.number:
        block.number = str(renderer.formula_counter)
        renderer.formula_counter += 1

    if block.formula_id:
        renderer.formula_refs[block.formula_id] = block.number

    # --- Вставка формулы ---
    add_formula_paragraph(
        renderer.doc,
        latex=block.latex,
        number=block.number,
        font_size_pt=DocumentSettings.FONT_SIZE_PT,
    )

    # --- Пояснение (где ...) ---
    if block.explanation:
        for raw_line in block.explanation.strip().split('\\n'):
            raw_line = raw_line.strip()
            if not raw_line:
                continue

            add_paragraph_with_inline_formulas(
                renderer.doc,
                text=raw_line,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent_cm=DocumentSettings.FIRST_LINE_INDENT_CM,
                font_size_pt=DocumentSettings.FONT_SIZE_PT,
            )

    renderer.doc.add_paragraph()  # отступ после формулы
    renderer._mark_content()

