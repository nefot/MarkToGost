# utils/document_helpers.py
"""Вспомогательные функции для работы с документами"""

import re
import os
import tempfile
import subprocess
from copy import deepcopy
from typing import Dict, Optional

from docx import Document
from docx.shared import Cm

from MarkToGost.config import DocumentSettings


def compute_image_width_cm(doc: Document, fraction: float = 0.70) -> Cm:
    """Вычисление ширины изображения на основе ширины страницы"""
    sec = doc.sections[0]
    try:
        page_width_cm = sec.page_width.cm
    except:
        page_width_cm = float(sec.page_width) / 360000.0

    left = sec.left_margin.cm
    right = sec.right_margin.cm
    avail = page_width_cm - left - right
    return Cm(avail * fraction)


def replace_image_refs(text: str, image_refs: Dict[str, int]) -> str:
    """Замена ссылок на изображения @img_id на их номера"""
    if not text or not image_refs:
        return text

    result = text
    for img_id, fig_num in image_refs.items():
        result = result.replace(f"@{img_id}", f"рис. {fig_num}")

    return result


def split_md_table_row(line: str) -> list:
    """Разбор строки markdown-таблицы"""
    s = line.strip()
    if s.startswith('|'):
        s = s[1:]
    if s.endswith('|'):
        s = s[:-1]

    cells = []
    buf = []
    escape = False

    for ch in s:
        if escape:
            buf.append(ch)
            escape = False
            continue
        if ch == '\\':
            escape = True
            continue
        if ch == '|':
            cells.append(''.join(buf).strip())
            buf = []
            continue
        buf.append(ch)

    cells.append(''.join(buf).strip())
    return cells


def is_md_table_separator(line: str) -> bool:
    """Проверка разделителя таблицы"""
    if '|' not in line:
        return False

    parts = split_md_table_row(line)
    if len(parts) < 2:
        return False

    for part in parts:
        p = part.replace(' ', '')
        if not re.fullmatch(r':?-{3,}:?', p):
            return False
    return True


def is_md_table_row(line: str) -> bool:
    """Проверка строки таблицы"""
    s = line.strip()
    return bool(s) and ('|' in s)


def normalize_table_caption(text: str) -> Optional[str]:
    """Извлечение названия таблицы из различных форматов"""
    if not text:
        return None

    text = re.sub(r'\s+', ' ', text).strip().rstrip('.')

    # Проверка различных форматов подписей
    patterns = [
        (r'^\s*<caption>\s*(.*?)\s*</caption>\s*$', re.I),  # HTML
        (r'^\s*(?:Таблица|Table)\s*\d*\s*[—:-]\s*(.+?)\s*$', re.I),  # Markdown-style
        (r'^\s*(?:Название таблицы|Caption)\s*[:\-]\s*(.+?)\s*$', re.I),  # Plain text
    ]

    for pattern_str, flags in patterns:
        m = re.match(pattern_str, text, flags)
        if m:
            title = m.group(1).strip().rstrip('.')
            return title or None

    return None


def append_docx_content(src_doc: Document, dst_doc: Document):
    """Копирует содержимое одного docx в другой"""
    for element in src_doc.element.body:
        dst_doc.element.body.append(deepcopy(element))


def render_formula_with_pandoc(latex: str) -> Document:
    """
    Генерирует DOCX с формулой через Pandoc или текстовый резервный вариант.
    
    Используется для конвертации TeX-формул в OMML (Office Open XML Math).
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        md_path = os.path.join(tmpdir, "formula.md")
        docx_path = os.path.join(tmpdir, "formula.docx")

        # ВАЖНО: display math syntax
        md_content = f"$${latex}$$"
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)

        try:
            subprocess.run([
                "pandoc",
                md_path,
                "-o",
                docx_path
            ], check=True, capture_output=True, timeout=10)

            return Document(docx_path)
        except (FileNotFoundError, subprocess.TimeoutExpired, subprocess.CalledProcessError):
            # Если pandoc не установлен или произошла ошибка,
            # создаем документ с текстом формулы
            doc = Document()
            p = doc.add_paragraph()
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            from docx.shared import Pt
            run = p.add_run(f"${latex}$")
            run.font.size = Pt(DocumentSettings.FONT_SIZE_PT)
            return doc

