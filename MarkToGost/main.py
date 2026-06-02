# -*- coding: utf-8 -*-
"""
MarkToGost: Конвертер Markdown → DOCX по ГОСТ 7.32-2001
Точка входа для пакетной обработки файлов

Использование:
    python -m MarkToGost.main                    # Обработка всех файлов из input/
    python -m MarkToGost.main input_file.md      # Обработка конкретного файла
    python -m MarkToGost.main --help             # Справка
"""

import os
import sys
import re
import argparse
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from MarkToGost.config import DocumentSettings
from MarkToGost.parser.blocks import Section, HeadingBlock
from MarkToGost.parser.markdown_parser import MarkdownParser
from MarkToGost.parser.metadata import extract_metadata

from MarkToGost.renderer.document_renderer import DocumentRenderer

from MarkToGost.utils.toc import (
    reset_heading_styles,
    get_toc_level,
    add_toc
)
from MarkToGost.utils.formatting import set_run_font, set_paragraph_formatting
from MarkToGost.utils.xml_helpers import add_page_number_centered

# Структура папок
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
# BASE_DIR это папка MarkToGost (содержит MarkToGost/main.py, MarkToGost/config.py и др.)
PROJECT_ROOT = BASE_DIR
INPUT_DIR = os.path.join(PROJECT_ROOT, "input")
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "output")
# Создаем папки если их нет
os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)


def apply_cli_overrides(args, settings):
    """Применяет CLI-аргументы к DocumentSettings"""
    mapping = {
        'font_name': 'FONT_NAME',
        'font_size': 'FONT_SIZE_PT',
        'line_spacing': 'LINE_SPACING',
        'first_line_indent': 'FIRST_LINE_INDENT_CM',
        'left_margin': 'LEFT_MARGIN_CM',
        'right_margin': 'RIGHT_MARGIN_CM',
        'top_margin': 'TOP_MARGIN_CM',
        'bottom_margin': 'BOTTOM_MARGIN_CM',
        'image_width': 'IMAGE_WIDTH_FRACTION',
        'caption_font_size': 'CAPTION_FONT_SIZE_PT',
        'caption_italic': 'CAPTION_ITALIC',
        'table_font_size': 'TABLE_FONT_SIZE_PT',
        'use_first_line_indent': 'USE_FIRST_LINE_INDENT',
    }
    for arg_name, setting_name in mapping.items():
        value = getattr(args, arg_name, None)
        if value is not None:
            setattr(settings, setting_name, value)


def create_document(md_text: str) -> Document:
    """
    Создание DOCX документа из Markdown с соблюдением ГОСТ 7.32-2001
    
    Args:
        md_text: Полный текст Markdown документа
        
    Returns:
        Document: DOCX документ (python-docx)
    """

    # Извлечение метаданных
    metadata = extract_metadata(md_text)

    # Очищаем служебные метаданные перед парсингом
    md_text_for_parsing = md_text
    md_text_for_parsing = re.sub(r'__?ФИО__?[:\-]?.*\n?', '', md_text_for_parsing)
    md_text_for_parsing = re.sub(r'__?Группа__?[:\-]?.*\n?', '', md_text_for_parsing)
    md_text_for_parsing = re.sub(r'\[//\]:\s*#\s*\w+\s*=.*\n?', '', md_text_for_parsing, flags=re.I)

    # Создание документа
    doc = Document()

    # Настройка страницы по ГОСТ
    for s in doc.sections:
        s.page_height = Cm(DocumentSettings.PAGE_HEIGHT_CM)
        s.page_width = Cm(DocumentSettings.PAGE_WIDTH_CM)
        s.left_margin = Cm(DocumentSettings.LEFT_MARGIN_CM)
        s.right_margin = Cm(DocumentSettings.RIGHT_MARGIN_CM)
        s.top_margin = Cm(DocumentSettings.TOP_MARGIN_CM)
        s.bottom_margin = Cm(DocumentSettings.BOTTOM_MARGIN_CM)

    # Настройка стилей документа
    normal = doc.styles['Normal']
    normal.font.name = DocumentSettings.FONT_NAME
    normal.font.size = Pt(DocumentSettings.FONT_SIZE_PT)
    normal.paragraph_format.first_line_indent = Cm(DocumentSettings.FIRST_LINE_INDENT_CM)
    normal.paragraph_format.line_spacing = DocumentSettings.LINE_SPACING
    reset_heading_styles(doc)

    # Блок ФИО (если не таблица подписей)
    if not metadata['is_table'] and (metadata['fio'] or metadata['group']):
        group_full = f"090302-{metadata['group']}" if metadata['group'] else ""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run1 = p.add_run(f"Выполнил студент: {metadata['fio']}\n")
        set_run_font(run1, size_pt=DocumentSettings.FONT_SIZE_PT, bold=True)
        run2 = p.add_run(f"Группа: {group_full}")
        set_run_font(run2, size_pt=DocumentSettings.FONT_SIZE_PT, bold=True)

    # Парсинг Markdown и сбор заголовков для оглавления
    parser = MarkdownParser(md_text_for_parsing)
    blocks = parser.parse()

    toc_entries = []

    def collect_headings(block):
        """Рекурсивный сбор всех заголовков для оглавления"""
        if isinstance(block, Section):
            for b in block.blocks:
                collect_headings(b)
        elif isinstance(block, HeadingBlock):
            toc_level = get_toc_level(block.text)
            toc_entries.append((toc_level, block.text))

    for block in blocks:
        collect_headings(block)

    # Рендеринг блоков в документ
    renderer = DocumentRenderer(doc, toc_entries, use_headings=metadata['use_headings'])
    for block in blocks:
        renderer.render_block(block)

    # Таблица подписей (если указано is_table: true)
    if metadata['is_table']:
        doc.add_page_break()

        from MarkToGost.utils.xml_helpers import set_table_borders, set_repeat_table_header
        from docx.enum.table import WD_TABLE_ALIGNMENT

        table = doc.add_table(rows=2, cols=2)
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        col_width = Cm(
            (DocumentSettings.PAGE_WIDTH_CM - DocumentSettings.LEFT_MARGIN_CM - DocumentSettings.RIGHT_MARGIN_CM) / 2.0
        )
        for col in table.columns:
            for cell in col.cells:
                cell.width = col_width

        table.cell(0, 0).text = 'Выполнил:'
        table.cell(0, 1).text = metadata['fio']
        table.cell(1, 0).text = 'Проверил:'
        table.cell(1, 1).text = metadata['teacher']

        # Форматирование таблицы
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:
                set_repeat_table_header(row)
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.style = doc.styles['Normal']
                    p.paragraph_format.first_line_indent = None
                    p.paragraph_format.left_indent = Cm(0)
                    p.paragraph_format.line_spacing = 1.5
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        set_run_font(run, size_pt=12)

        set_table_borders(table)

    # Добавление номеров страниц в колонтитул
    if metadata['numerate']:
        add_page_number_centered(doc)

    return doc


def process_md_file(input_path: str, output_path: Optional[str] = None) -> bool:
    """
    Обработка одного Markdown файла: парсинг и создание DOCX
    
    Args:
        input_path: Путь к .md файлу (абсолютный или относительный к INPUT_DIR)
        output_path: Путь для сохранения DOCX (если None, используется OUTPUT_DIR)
        
    Returns:
        bool: True если успешно, False если ошибка
    """
    # Нормализуем входной путь
    if not os.path.isabs(input_path):
        input_full_path = os.path.join(INPUT_DIR, input_path)
    else:
        input_full_path = input_path

    # Проверяем что файл существует
    if not os.path.exists(input_full_path):
        print(f"❌ Файл не найден: {input_full_path}")
        return False

    # Определяем путь выхода
    if output_path is None:
        filename = os.path.basename(input_full_path)
        name_wo_ext = os.path.splitext(filename)[0]
        output_full_path = os.path.join(OUTPUT_DIR, f"{name_wo_ext}.docx")
    else:
        if not os.path.isabs(output_path):
            output_full_path = os.path.join(OUTPUT_DIR, output_path)
        else:
            output_full_path = output_path
        # Создаем директорию если нужно
        os.makedirs(os.path.dirname(output_full_path), exist_ok=True)

    try:
        with open(input_full_path, "r", encoding="utf-8") as f:
            md_text = f.read()

        doc = create_document(md_text)
        doc.save(output_full_path)

        rel_input = os.path.relpath(input_full_path, PROJECT_ROOT)
        rel_output = os.path.relpath(output_full_path, PROJECT_ROOT)
        print(f"✅ {rel_input} → {rel_output}")
        return True

    except Exception as e:
        print(f"❌ Ошибка при обработке {os.path.basename(input_full_path)}: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="MarkToGost: Конвертер Markdown → DOCX по ГОСТ 7.32-2001",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
        Примеры использования:
          python -m MarkToGost.main
                Обработает все .md файлы из папки 'input/'
        
          python -m MarkToGost.main file.md
                Обработает конкретный файл 'input/file.md'
        
          python -m MarkToGost.main file.md --output custom_output.docx
                Обработает 'input/file.md' и сохранит результат как 'output/custom_output.docx'
        
          python -m MarkToGost.main file.md --font-size 12 --left-margin 2.5
                Обработает файл с нестандартными параметрами документа
        
        Структура папок:
          input/      - входные Markdown файлы
          output/     - результаты конвертации (DOCX файлы)
        """
    )

    parser.add_argument(
        "file",
        nargs="?",
        default=None,
        help="Конкретный MD файл для обработки (из папки input/)"
    )

    parser.add_argument(
        "--output", "-o",
        default=None,
        help="Имя выходного DOCX файла (будет сохранен в папку output/)"
    )

    # --- Настройки документа ---
    doc_group = parser.add_argument_group('Настройки документа')
    doc_group.add_argument('--font-name', default=None, help=f'Шрифт (по умол.: {DocumentSettings.FONT_NAME})')
    doc_group.add_argument('--font-size', default=None, type=float,
                           help=f'Размер шрифта в pt (по умол.: {DocumentSettings.FONT_SIZE_PT})')
    doc_group.add_argument('--line-spacing', default=None, type=float,
                           help=f'Межстрочный интервал (по умол.: {DocumentSettings.LINE_SPACING})')
    doc_group.add_argument('--first-line-indent', default=None, type=float,
                           help=f'Отступ первой строки в см (по умол.: {DocumentSettings.FIRST_LINE_INDENT_CM})')
    doc_group.add_argument('--left-margin', default=None, type=float,
                           help=f'Левое поле в см (по умол.: {DocumentSettings.LEFT_MARGIN_CM})')
    doc_group.add_argument('--right-margin', default=None, type=float,
                           help=f'Правое поле в см (по умол.: {DocumentSettings.RIGHT_MARGIN_CM})')
    doc_group.add_argument('--top-margin', default=None, type=float,
                           help=f'Верхнее поле в см (по умол.: {DocumentSettings.TOP_MARGIN_CM})')
    doc_group.add_argument('--bottom-margin', default=None, type=float,
                           help=f'Нижнее поле в см (по умол.: {DocumentSettings.BOTTOM_MARGIN_CM})')
    doc_group.add_argument('--image-width', default=None, type=float,
                           help=f'Ширина изображений — доля от страницы (по умол.: {DocumentSettings.IMAGE_WIDTH_FRACTION})')
    doc_group.add_argument('--caption-font-size', default=None, type=float,
                           help=f'Размер шрифта подписей в pt (по умол.: {DocumentSettings.CAPTION_FONT_SIZE_PT})')
    doc_group.add_argument('--caption-italic', default=None, type=lambda x: x.lower() == 'true',
                           help=f'Курсив подписей: true/false (по умол.: {DocumentSettings.CAPTION_ITALIC})')
    doc_group.add_argument('--table-font-size', default=None, type=float,
                           help=f'Размер шрифта таблиц в pt (по умол.: {DocumentSettings.TABLE_FONT_SIZE_PT})')
    doc_group.add_argument('--use-first-line-indent', default=None, type=lambda x: x.lower() == 'true',
                           help=f'Красная строка: true/false (по умол.: {DocumentSettings.USE_FIRST_LINE_INDENT})')

    args = parser.parse_args()

    # Применяем CLI-параметры к конфигу ДО обработки файлов
    apply_cli_overrides(args, DocumentSettings)

    # Если указан конкретный файл
    if args.file:
        print(f"📄 Обработка файла: {args.file}")
        success = process_md_file(args.file, args.output)
        sys.exit(0 if success else 1)

    # Иначе обрабатываем все файлы из input/
    print(f"📂 Папка входных файлов: {INPUT_DIR}")
    print(f"📂 Папка для сохранения: {OUTPUT_DIR}")

    all_md = [
        f for f in os.listdir(INPUT_DIR)
        if f.lower().endswith(".md") and not f.startswith(".")
    ]

    if not all_md:
        print(f"⚠️ Не найдено ни одного Markdown-файла в {INPUT_DIR}")
        sys.exit(1)

    print(f"🔍 Найдено файлов: {len(all_md)}\n")

    success_count = 0
    for file in all_md:
        if process_md_file(os.path.join(INPUT_DIR, file)):
            success_count += 1

    print(f"\n🎉 Обработано: {success_count}/{len(all_md)} файлов")
    sys.exit(0 if success_count == len(all_md) else 1)
