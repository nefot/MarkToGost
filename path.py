"""
ПАТЧ: Исправление парсинга и рендеринга формул
===============================================

ПРОБЛЕМА 1: Строки "где $F$ — сила..." после $$ не попадали в FormulaBlock.explanation
ПРОБЛЕМА 2: Inline math $...$ в explanation рендерился как сырой текст с долларами
ПРОБЛЕМА 3: Пояснения слипались в одну строку без пробелов

РЕШЕНИЕ:
- _parse_formula_block() теперь собирает строки "где..." в explanation
- Новая функция convert_inline_math() убирает $ и форматирует inline-math
- _render_formula_block() корректно рендерит каждую строку explanation
"""

import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ================================
# 1. КОНВЕРТАЦИЯ INLINE MATH $...$
# ================================

def convert_inline_math(text: str) -> str:
    """
    Конвертирует inline math $...$ в читаемый текст.
    Убирает знаки доллара, оставляя содержимое.

    Примеры:
        $F$ → F
        $m_1$ → m₁  (нижний индекс через юникод)
        $6.674 \\times 10^{-11}$ → 6.674×10⁻¹¹
        $E_k$ → Eₖ
    """
    if not text:
        return text

    # Словари для конвертации индексов в юникод
    SUBSCRIPT_MAP = str.maketrans("0123456789aeinoruvxhklmnpst", "₀₁₂₃₄₅₆₇₈₉ₐₑᵢₙₒᵣᵤᵥₓₕₖₗₘₙₚₛₜ")
    SUPERSCRIPT_MAP = str.maketrans("0123456789+-=()ni", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿⁱ")

    def process_math(m: str) -> str:
        """Обрабатывает содержимое между знаками $"""
        # Замены LaTeX команд
        m = re.sub(r'\s*\\times\s*', '×', m)
        m = m.replace(r'\cdot', '·')
        m = m.replace(r'\pm', '±')
        m = m.replace(r'\leq', '≤')
        m = m.replace(r'\geq', '≥')
        m = m.replace(r'\neq', '≠')
        m = m.replace(r'\alpha', 'α')
        m = m.replace(r'\beta', 'β')
        m = m.replace(r'\gamma', 'γ')
        m = m.replace(r'\delta', 'δ')
        m = m.replace(r'\Delta', 'Δ')
        m = m.replace(r'\sigma', 'σ')
        m = m.replace(r'\rho', 'ρ')
        m = m.replace(r'\mu', 'μ')
        m = m.replace(r'\lambda', 'λ')
        m = m.replace(r'\omega', 'ω')
        m = m.replace(r'\Omega', 'Ω')
        m = m.replace(r'\pi', 'π')
        m = m.replace(r'\infty', '∞')
        m = m.replace(r'\frac', '')

        # Нижние индексы _x или _{xxx}
        def replace_sub(match):
            content = match.group(1) or match.group(2)
            return content.translate(SUBSCRIPT_MAP)

        m = re.sub(r'_\{([^}]+)\}|_([^{}\s])', replace_sub, m)

        # Верхние индексы ^x или ^{xxx}
        def replace_sup(match):
            content = match.group(1) or match.group(2)
            return content.translate(SUPERSCRIPT_MAP)

        m = re.sub(r'\^\{([^}]+)\}|\^([^{}\s])', replace_sup, m)

        # Убираем оставшиеся фигурные скобки
        m = m.replace('{', '').replace('}', '')

        return m.strip()

    # Заменяем $...$
    result = re.sub(r'\$([^$]+)\$', lambda match: process_math(match.group(1)), text)
    return result


# ================================
# 2. ИСПРАВЛЕННЫЙ ПАРСЕР ФОРМУЛ
# ================================

def parse_formula_block_fixed(self):
    """
    ЗАМЕНА для MarkdownParser._parse_formula_block()

    Собирает:
    - Саму формулу LaTeX между $$...$$
    - Строки пояснения "где ..." идущие сразу после формулы
    """
    line = self.lines[self.index].strip()
    line = line[2:].strip()  # Убираем начальные $$

    buffer = []

    # Собираем тело формулы до закрывающего $$
    while True:
        if "$$" in line:
            before, _, _ = line.partition("$$")
            before = before.strip()
            if before:
                buffer.append(before)
            break
        else:
            if line:
                buffer.append(line)
            self.index += 1
            if self.index >= len(self.lines):
                break
            line = self.lines[self.index].strip()

    self.index += 1  # Пропускаем строку с закрывающим $$

    formula_text = "\n".join(buffer).strip()

    # ============================
    # НОВОЕ: Собираем explanation
    # ============================
    explanation_lines = []

    while self.index < len(self.lines):
        peek = self.lines[self.index]
        peek_stripped = peek.strip()

        # Останавливаемся если:
        if not peek_stripped:
            # Пустая строка — проверяем следующую
            # Если после пустой идёт ещё пояснение — продолжаем
            if self.index + 1 < len(self.lines):
                next_line = self.lines[self.index + 1].strip()
                # Если следующая строка начинается с $ (переменная) или где — это продолжение
                if next_line.startswith('$') or next_line.lower().startswith('где'):
                    self.index += 1
                    continue
            # Иначе — конец пояснения
            break

        if peek_stripped.startswith('#'):  # Новый заголовок
            break
        if peek_stripped.startswith('$$'):  # Новая формула
            break
        if peek_stripped.startswith('!['):  # Изображение
            break
        if peek_stripped.startswith('[//]:'):  # Раздел
            break

        # Это строка пояснения
        explanation_lines.append(peek_stripped)
        self.index += 1

    # Объединяем пояснение
    explanation = "\n".join(explanation_lines).strip() if explanation_lines else None

    from dataclasses import fields
    # Импортируем FormulaBlock из основного модуля
    try:
        from __main__ import FormulaBlock
    except ImportError:
        # Используем локальное определение для тестирования
        pass

    return formula_text, explanation


# ================================
# 3. ИСПРАВЛЕННЫЙ РЕНДЕРЕР ФОРМУЛ
# ================================

def render_formula_block_fixed(self, block):
    """
    ЗАМЕНА для DocumentRenderer._render_formula_block()

    Исправления:
    - explanation конвертируется через convert_inline_math()
    - Каждая строка explanation рендерится отдельно
    - Правильное форматирование по ГОСТ (отступ, выравнивание)
    """
    from copy import deepcopy
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # --- Нумерация ---
    if not block.number:
        block.number = str(self.formula_counter)
        self.formula_counter += 1

    if block.formula_id:
        self.formula_refs[block.formula_id] = block.number

    # --- Получаем формулу через Pandoc ---
    formula_doc = render_formula_with_pandoc(block.latex)

    # --- Создаём таблицу 1x2: [формула] [номер] ---
    table = self.doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    sec = self.doc.sections[0]
    usable_width = sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm

    table.columns[0].width = Cm(usable_width * 0.85)
    table.columns[1].width = Cm(usable_width * 0.15)

    # Левая ячейка — формула
    left_cell = table.cell(0, 0)
    left_cell.text = ""
    for el in formula_doc.element.body:
        left_cell._element.append(deepcopy(el))
    for p in left_cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Правая ячейка — номер
    right_cell = table.cell(0, 1)
    right_cell.text = ""
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"({block.number})")
    set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT)

    # Убираем границы таблицы
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tblPr.append(borders)

    # --- Пояснение (ИСПРАВЛЕНО) ---
    if block.explanation:
        lines = block.explanation.strip().split('\n')

        for i, raw_line in enumerate(lines):
            raw_line = raw_line.strip()
            if not raw_line:
                continue

            # Конвертируем inline math $...$ → читаемый текст
            line_text = convert_inline_math(raw_line)

            exp_p = self.doc.add_paragraph()
            set_paragraph_formatting(
                exp_p,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=Cm(DocumentSettings.FIRST_LINE_INDENT_CM),
                line_spacing=DocumentSettings.LINE_SPACING,
                space_before=0,
                space_after=0
            )

            run = exp_p.add_run(line_text)
            set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT)

    # Пустой абзац после блока
    self.doc.add_paragraph()
    self._mark_content()


# ================================
# 4. КАК ПРИМЕНИТЬ ПАТЧ
# ================================

PATCH_INSTRUCTIONS = """
ИНСТРУКЦИЯ ПО ПРИМЕНЕНИЮ ПАТЧА
================================

1. ПАРСЕР (MarkdownParser._parse_formula_block):

Замените метод на следующий код:

    def _parse_formula_block(self):
        line = self.lines[self.index].strip()
        line = line[2:].strip()  # Убираем начальные $$

        buffer = []
        while True:
            if "$$" in line:
                before, _, _ = line.partition("$$")
                before = before.strip()
                if before:
                    buffer.append(before)
                break
            else:
                if line:
                    buffer.append(line)
                self.index += 1
                if self.index >= len(self.lines):
                    break
                line = self.lines[self.index].strip()

        self.index += 1  # Пропускаем строку с $$

        formula_text = "\\n".join(buffer).strip()

        # Собираем строки пояснения "где ..."
        explanation_lines = []
        while self.index < len(self.lines):
            peek = self.lines[self.index].strip()

            if not peek:
                # Пустая строка — смотрим вперёд
                if self.index + 1 < len(self.lines):
                    next_line = self.lines[self.index + 1].strip()
                    if next_line.startswith('$') or next_line.lower().startswith('где'):
                        self.index += 1
                        continue
                break

            # Стоп-условия
            if (peek.startswith('#') or peek.startswith('$$') or 
                peek.startswith('![') or peek.startswith('[//]:')):
                break

            explanation_lines.append(peek)
            self.index += 1

        explanation = "\\n".join(explanation_lines).strip() if explanation_lines else None

        return FormulaBlock(latex=formula_text, explanation=explanation)


2. ДОБАВЬТЕ функцию convert_inline_math() в начало файла (после импортов).
   Полный код функции — в этом файле выше.


3. РЕНДЕРЕР (DocumentRenderer._render_formula_block):

В методе _render_formula_block замените блок рендеринга explanation:

    # СТАРЫЙ КОД (удалить):
    if block.explanation:
        exp_p = self.doc.add_paragraph()
        ...
        exp_p.add_run("где ")
        exp_p.add_run(block.explanation)

    # НОВЫЙ КОД:
    if block.explanation:
        lines = block.explanation.strip().split('\\n')
        for raw_line in lines:
            raw_line = raw_line.strip()
            if not raw_line:
                continue
            line_text = convert_inline_math(raw_line)
            exp_p = self.doc.add_paragraph()
            set_paragraph_formatting(
                exp_p,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=Cm(DocumentSettings.FIRST_LINE_INDENT_CM),
                line_spacing=DocumentSettings.LINE_SPACING,
                space_before=0,
                space_after=0
            )
            run = exp_p.add_run(line_text)
            set_run_font(run, size_pt=DocumentSettings.FONT_SIZE_PT)
"""

if __name__ == "__main__":
    # Тест convert_inline_math
    test_cases = [
        ("$E$ — энергия, Дж; $m$ — масса, кг; $c$ — скорость света, м/с.",
         "E — энергия, Дж; m — масса, кг; c — скорость света, м/с."),
        ("$Q$ — количество теплоты, Дж;",
         "Q — количество теплоты, Дж;"),
        ("$G$ — гравитационная постоянная, $6.674 \\times 10^{-11}$ м³/(кг·с²);",
         "G — гравитационная постоянная, 6.674×10⁻¹¹ м³/(кг·с²);"),
        ("$m_1, m_2$ — массы тел, кг;",
         "m₁, m₂ — массы тел, кг;"),
        ("$E_k$ — кинетическая энергия, Дж;",
         "Eₖ — кинетическая энергия, Дж;"),
    ]

    print("=== ТЕСТ convert_inline_math ===\n")
    all_ok = True
    for input_text, expected in test_cases:
        result = convert_inline_math(input_text)
        status = "✅" if result == expected else "⚠️ "
        if result != expected:
            all_ok = False
        print(f"{status} Вход:    {input_text}")
        print(f"   Результат: {result}")
        if result != expected:
            print(f"   Ожидалось: {expected}")
        print()

    print("✅ Все тесты пройдены!" if all_ok else "⚠️  Некоторые тесты не совпали — проверьте вывод выше")
    print("\n" + PATCH_INSTRUCTIONS)